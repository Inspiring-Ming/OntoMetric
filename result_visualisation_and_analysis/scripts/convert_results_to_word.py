#!/usr/bin/env python3
"""
Convert ESG Metric Extraction Experiments Results from Markdown to Word
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path
import re

def parse_markdown_table(lines):
    """Parse markdown table into rows"""
    rows = []
    for line in lines:
        if '|' in line and not line.strip().startswith('|---'):
            cells = [cell.strip() for cell in line.split('|')]
            cells = [c for c in cells if c]  # Remove empty cells
            if cells:
                rows.append(cells)
    return rows

def add_formatted_text(paragraph, text):
    """Add text with markdown formatting (bold, italic, code)"""
    # Handle bold
    parts = re.split(r'\*\*(.+?)\*\*', text)
    for i, part in enumerate(parts):
        if i % 2 == 0:
            # Handle italic
            italic_parts = re.split(r'\*(.+?)\*', part)
            for j, ipart in enumerate(italic_parts):
                if j % 2 == 0:
                    # Handle code
                    code_parts = re.split(r'`(.+?)`', ipart)
                    for k, cpart in enumerate(code_parts):
                        if k % 2 == 0:
                            if cpart:
                                paragraph.add_run(cpart)
                        else:
                            run = paragraph.add_run(cpart)
                            run.font.name = 'Courier New'
                            run.font.size = Pt(9)
                else:
                    run = paragraph.add_run(ipart)
                    run.italic = True
        else:
            run = paragraph.add_run(part)
            run.bold = True

def convert_markdown_to_word(md_path, output_path):
    """Convert markdown file to Word document"""
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        # Title (# heading)
        if line.startswith('# '):
            title = line[2:].strip()
            p = doc.add_paragraph(title)
            p.style = 'Heading 1'
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Heading 2 (##)
        elif line.startswith('## '):
            heading = line[3:].strip()
            p = doc.add_paragraph(heading)
            p.style = 'Heading 2'

        # Heading 3 (###)
        elif line.startswith('### '):
            heading = line[4:].strip()
            p = doc.add_paragraph(heading)
            p.style = 'Heading 3'

        # Horizontal rule
        elif line.startswith('---'):
            doc.add_paragraph()  # Just add spacing

        # Table
        elif line.startswith('|') and i + 1 < len(lines) and '|---' in lines[i + 1]:
            # Collect table lines
            table_lines = [line]
            j = i + 1
            while j < len(lines) and lines[j].strip().startswith('|'):
                table_lines.append(lines[j].rstrip())
                j += 1

            rows = parse_markdown_table(table_lines)
            if len(rows) >= 2:
                # Create table (skip separator row)
                header = rows[0]
                data_rows = [r for r in rows[1:] if r and len(r) == len(header)]

                if data_rows:
                    table = doc.add_table(rows=len(data_rows) + 1, cols=len(header))
                    table.style = 'Light Grid Accent 1'

                    # Header row
                    for col_idx, cell_text in enumerate(header):
                        cell = table.rows[0].cells[col_idx]
                        cell.text = cell_text
                        cell.paragraphs[0].runs[0].bold = True
                        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                    # Data rows
                    for row_idx, row_data in enumerate(data_rows):
                        for col_idx, cell_text in enumerate(row_data):
                            if col_idx < len(table.rows[row_idx + 1].cells):
                                cell = table.rows[row_idx + 1].cells[col_idx]
                                # Remove markdown symbols
                                clean_text = cell_text.replace('**', '').replace('⭐', '⭐').replace('⚠️', '⚠️')
                                cell.text = clean_text

                doc.add_paragraph()  # Add spacing after table

            i = j - 1

        # Bold text paragraph
        elif line.startswith('**') and line.endswith('**'):
            text = line.strip('*').strip()
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.bold = True

        # Bullet list
        elif line.strip().startswith('- '):
            text = line.strip()[2:]
            p = doc.add_paragraph(style='List Bullet')
            add_formatted_text(p, text)

        # Numbered list or regular paragraph starting with number
        elif re.match(r'^\d+\.', line.strip()):
            text = re.sub(r'^\d+\.\s*', '', line.strip())
            if not text:  # Empty numbered line
                i += 1
                continue
            # Check if it's part of a list or just a numbered item
            p = doc.add_paragraph(style='List Number')
            add_formatted_text(p, text)

        # Regular paragraph
        elif line.strip():
            p = doc.add_paragraph()
            add_formatted_text(p, line.strip())

        # Empty line - skip
        else:
            pass

        i += 1

    # Save document
    doc.save(output_path)
    print(f"✅ Converted to Word: {output_path}")


if __name__ == '__main__':
    md_file = Path(__file__).parent.parent / "ESG_Metric_Extraction_Experiments_Results.md"
    output_file = Path(__file__).parent.parent / "ESG_Metric_Extraction_Experiments_Results.docx"

    print("=" * 80)
    print("CONVERTING MARKDOWN TO WORD")
    print("=" * 80)
    convert_markdown_to_word(md_file, output_file)
    print("=" * 80)
