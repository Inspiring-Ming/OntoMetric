from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

def parse_markdown_table(lines, start_idx):
    """Parse a markdown table starting at start_idx"""
    if start_idx >= len(lines):
        return None, start_idx

    # Find table header
    header_idx = start_idx
    while header_idx < len(lines) and not lines[header_idx].strip().startswith('|'):
        header_idx += 1

    if header_idx >= len(lines):
        return None, start_idx

    # Parse header
    header_line = lines[header_idx].strip()
    headers = [h.strip() for h in header_line.split('|')[1:-1]]

    # Skip separator line
    separator_idx = header_idx + 1
    if separator_idx >= len(lines) or not lines[separator_idx].strip().startswith('|'):
        return None, start_idx

    # Parse rows
    rows = []
    row_idx = separator_idx + 1
    while row_idx < len(lines) and lines[row_idx].strip().startswith('|'):
        row_line = lines[row_idx].strip()
        cells = [c.strip() for c in row_line.split('|')[1:-1]]
        rows.append(cells)
        row_idx += 1

    return {'headers': headers, 'rows': rows}, row_idx

def add_formatted_text(paragraph, text, bold=False, size=None):
    """Add formatted text to a paragraph"""
    run = paragraph.add_run(text)
    if bold:
        run.bold = True
    if size:
        run.font.size = Pt(size)
    return run

def add_markdown_table_to_doc(doc, table_data):
    """Add a markdown table to the Word document"""
    if not table_data or not table_data['rows']:
        return

    headers = table_data['headers']
    rows = table_data['rows']

    # Create table
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.style = 'Light Grid Accent 1'

    # Add headers
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    # Add data rows
    for row_idx, row in enumerate(rows):
        for col_idx, cell_value in enumerate(row):
            table.rows[row_idx + 1].cells[col_idx].text = cell_value

    doc.add_paragraph()  # Add spacing after table

def convert_markdown_to_word(md_file, docx_file):
    """Convert markdown file to Word document"""
    # Read markdown file
    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    # Create Word document
    doc = Document()

    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        # Title (# )
        if line.startswith('# ') and not line.startswith('## '):
            title = line[2:].strip()
            p = doc.add_heading(title, level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Section heading (## )
        elif line.startswith('## ') and not line.startswith('### '):
            heading = line[3:].strip()
            doc.add_heading(heading, level=2)

        # Subsection heading (### )
        elif line.startswith('### '):
            subheading = line[4:].strip()
            doc.add_heading(subheading, level=3)

        # Subsubsection heading (#### )
        elif line.startswith('#### '):
            subsubheading = line[5:].strip()
            doc.add_heading(subsubheading, level=4)

        # Horizontal rule
        elif line.strip() == '---':
            doc.add_paragraph('_' * 80)

        # Table
        elif line.strip().startswith('|'):
            table_data, next_idx = parse_markdown_table(lines, i)
            if table_data:
                add_markdown_table_to_doc(doc, table_data)
                i = next_idx - 1

        # Bold text pattern **text:**
        elif line.strip().startswith('**') and ':**' in line:
            p = doc.add_paragraph()
            add_formatted_text(p, line.strip().replace('**', ''), bold=True)

        # Regular paragraph
        elif line.strip() and not line.strip().startswith('#'):
            doc.add_paragraph(line.strip())

        i += 1

    # Save document
    doc.save(docx_file)
    print(f"✓ Created: {docx_file}")

def main():
    md_file = './result_visualisation_and_analysis/All_Documents_Results.md'
    docx_file = './result_visualisation_and_analysis/All_Documents_Results.docx'

    print("Converting All_Documents_Results.md to Word format with detailed entity tables...")
    convert_markdown_to_word(md_file, docx_file)
    print("\nConversion complete!")

if __name__ == '__main__':
    main()
